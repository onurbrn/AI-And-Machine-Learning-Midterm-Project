from docx import Document
import openai
import os

openai.api_key = "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"

folder_path = "C:\\Users\\copro\\OneDrive\\Masaüstü\\Machine"

def read_docx_file(file_path):
    """
    Reads the content of a .docx file and returns the text as a string.

    Args:
        file_path (str): Path to the .docx file.
    
    Returns:
        str: Extracted text from the document.
    """
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])

def analyze_document(file_name, content):
    """
    Analyzes the document content using GPT-4 and provides a detailed reasoning 
    and summary with a structured decision-making approach.

    Args:
        file_name (str): Name of the document being analyzed.
        content (str): Text content of the document.
    
    Returns:
        str: Analysis result with reasoning and summary.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4o",  
        messages=[
            {"role": "system", "content": (
                "You are a highly experienced forensic investigator tasked with analyzing written evidence. "
                "Follow these steps for detailed reasoning and analysis:\n"
                "1. Identify the key elements of the document (e.g., evidence descriptions, timelines, suspect details).\n"
                "2. Assess how each element contributes to the investigation.\n"
                "3. Highlight potential links, contradictions, or unique findings in the document.\n"
                "4. Provide a detailed reasoning process within <Reasoning> tags.\n"
                "5. Summarize the key findings within <Summary> tags."
            )},
            {"role": "user", "content": f"Analyze the following document ({file_name}):\n{content}"}
        ],
        temperature=0.5,  
        max_tokens=1500
    )
    return response['choices'][0]['message']['content']

def save_combined_analysis_to_docx(output_file, analyses):
    """
    Saves all analyses to a single .docx file.

    Args:
        output_file (str): Name of the output .docx file.
        analyses (list): List of analyses with document names and content.
    """
    doc = Document()
    doc.add_heading("Combined Analysis Report", level=1)
    
    for analysis in analyses:
        doc.add_heading(f"Analysis for {analysis['file_name']}", level=2)
        doc.add_paragraph(analysis['content'])
    
    doc.save(output_file)
    print(f"All analyses saved to {output_file}")

def process_all_documents(folder_path, output_file):
    """
    Processes all .docx files in a folder, analyzes their content, 
    and combines the results into a single .docx file.

    Args:
        folder_path (str): Path to the folder containing .docx files.
        output_file (str): Name of the output .docx file.
    """
    analyses = []
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            file_path = os.path.join(folder_path, file_name)
            print(f"Reading and analyzing: {file_name}")
            content = read_docx_file(file_path)
            analysis = analyze_document(file_name, content)
            analyses.append({"file_name": file_name, "content": analysis})
    
    save_combined_analysis_to_docx(output_file, analyses)

output_file = "Combined_Analysis_Report2.docx"
process_all_documents(folder_path, output_file)