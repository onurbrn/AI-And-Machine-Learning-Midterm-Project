import os
import openai
from docx import Document

openai.api_key = "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"

def transcribe_audio(audio_file):
    """
    Transcribes an MP3 audio file to text using OpenAI's Whisper model.
    
    Args:
        audio_file (str): Path to the MP3 file to transcribe
        
    Returns:
        str: The transcribed text
    """
    try:
        with open(audio_file, "rb") as audio:
            transcription = openai.Audio.transcribe(
                model="whisper-1",
                file=audio
            )
            
        return transcription['text']
        
    except Exception as e:
        print(f"An error occurred during transcription for {audio_file}: {str(e)}")
        return None

audio_folder = "C:\\Users\\copro\\OneDrive\\Masaüstü\\machinevoice"  
output_file = "transcriptions.docx"  

audio_files = [f for f in os.listdir(audio_folder) if f.endswith(".mp3")]

doc = Document()
doc.add_heading('Transcriptions', 0)

for audio_file in audio_files:
    file_path = os.path.join(audio_folder, audio_file)
    print(f"Processing file: {file_path}")
    
    transcription = transcribe_audio(file_path)
    
    if transcription:
        doc.add_paragraph(f"File: {audio_file}")
        doc.add_paragraph(transcription)
        doc.add_paragraph()  

doc.save(output_file)

print(f"\nAll transcriptions saved to {output_file}")
